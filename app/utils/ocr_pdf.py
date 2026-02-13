# -*- coding: utf-8 -*-
"""
Utilitário para OCR (Reconhecimento Óptico de Caracteres) em PDFs.
Usa Tesseract OCR para extrair texto de PDFs escaneados ou imagens.
"""
import os
import sys
from io import BytesIO
import fitz  # PyMuPDF
from PIL import Image
import pytesseract

# ==================== CONFIGURAÇÃO DO TESSERACT ====================

def get_project_root():
    """Retorna o diretório raiz do projeto."""
    # app/utils/ocr_pdf.py -> voltar 2 níveis
    current_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.dirname(os.path.dirname(current_dir))


def find_tesseract():
    """
    Procura o executável do Tesseract em múltiplos locais.
    Prioridade:
    1. Pasta 'tesseract' dentro do projeto
    2. Caminhos comuns de instalação no Windows
    3. PATH do sistema (padrão do pytesseract)
    
    Returns:
        tuple: (caminho_tesseract, caminho_tessdata) ou (None, None)
    """
    project_root = get_project_root()
    
    # Lista de possíveis locais do Tesseract
    possible_paths = [
        # Pasta local no projeto (prioridade máxima)
        os.path.join(project_root, 'tesseract'),
        os.path.join(project_root, 'tesseract', 'Tesseract-OCR'),
        # Caminhos comuns de instalação no Windows
        r'C:\Program Files\Tesseract-OCR',
        r'C:\Program Files (x86)\Tesseract-OCR',
        os.path.expanduser(r'~\AppData\Local\Tesseract-OCR'),
        os.path.expanduser(r'~\AppData\Local\Programs\Tesseract-OCR'),
        # Caminho relativo ao executável (para versão compilada)
        os.path.join(os.path.dirname(sys.executable), 'tesseract'),
        os.path.join(os.path.dirname(sys.executable), 'Tesseract-OCR'),
    ]
    
    for path in possible_paths:
        tesseract_exe = os.path.join(path, 'tesseract.exe')
        if os.path.isfile(tesseract_exe):
            tessdata_dir = os.path.join(path, 'tessdata')
            return tesseract_exe, tessdata_dir
    
    return None, None


def configure_tesseract():
    """
    Configura o pytesseract para usar o Tesseract encontrado.
    Retorna informações sobre a configuração.
    """
    tesseract_path, tessdata_path = find_tesseract()
    
    if tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        if tessdata_path and os.path.isdir(tessdata_path):
            os.environ['TESSDATA_PREFIX'] = tessdata_path
        return {
            'configured': True,
            'tesseract_path': tesseract_path,
            'tessdata_path': tessdata_path
        }
    
    return {'configured': False}


# Configurar Tesseract na inicialização do módulo
_tesseract_config = configure_tesseract()

# Idiomas disponíveis
LANGUAGES = {
    'por': 'Português',
    'eng': 'Inglês',
    'spa': 'Espanhol',
    'fra': 'Francês',
    'deu': 'Alemão',
    'ita': 'Italiano',
    'por+eng': 'Português + Inglês'
}


def check_tesseract_installed():
    """Verifica se o Tesseract está instalado e configurado."""
    global _tesseract_config
    
    # Tentar reconfigurar se não estava configurado
    if not _tesseract_config.get('configured'):
        _tesseract_config = configure_tesseract()
    
    try:
        version = pytesseract.get_tesseract_version()
        project_root = get_project_root()
        local_path = os.path.join(project_root, 'tesseract')
        
        return {
            'installed': True,
            'version': str(version),
            'tesseract_path': _tesseract_config.get('tesseract_path', 'PATH do sistema'),
            'tessdata_path': _tesseract_config.get('tessdata_path'),
            'local_folder': local_path,
            'using_local': _tesseract_config.get('tesseract_path', '').startswith(project_root) if _tesseract_config.get('tesseract_path') else False
        }
    except Exception as e:
        project_root = get_project_root()
        local_path = os.path.join(project_root, 'tesseract')
        
        return {
            'installed': False,
            'error': str(e),
            'local_folder': local_path,
            'instructions': f'Copie a pasta Tesseract-OCR para: {local_path}'
        }


def get_available_languages():
    """Retorna os idiomas disponíveis no Tesseract instalado."""
    try:
        langs = pytesseract.get_languages()
        available = {}
        for code, name in LANGUAGES.items():
            # Para idiomas combinados, verificar ambos
            if '+' in code:
                codes = code.split('+')
                if all(c in langs for c in codes):
                    available[code] = name
            elif code in langs:
                available[code] = name
        return available
    except Exception:
        return LANGUAGES  # Retorna todos se não conseguir verificar


def pdf_to_images(pdf_content, dpi=300):
    """
    Converte páginas de um PDF em imagens.
    
    Args:
        pdf_content: Conteúdo binário do PDF
        dpi: Resolução das imagens (padrão 300)
    
    Returns:
        Lista de imagens PIL
    """
    doc = fitz.open(stream=pdf_content, filetype="pdf")
    images = []
    
    zoom = dpi / 72  # 72 é a resolução padrão do PDF
    matrix = fitz.Matrix(zoom, zoom)
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=matrix)
        
        # Converter para PIL Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    
    doc.close()
    return images


def ocr_image(image, lang='por'):
    """
    Executa OCR em uma imagem.
    
    Args:
        image: Imagem PIL
        lang: Código do idioma (ex: 'por', 'eng', 'por+eng')
    
    Returns:
        Texto extraído
    """
    # Configurações para melhor qualidade
    config = '--oem 3 --psm 3'  # LSTM engine, modo automático de segmentação
    
    text = pytesseract.image_to_string(image, lang=lang, config=config)
    return text


def ocr_pdf_to_text(file, lang='por', dpi=300):
    """
    Extrai texto de um PDF usando OCR.
    
    Args:
        file: Arquivo PDF
        lang: Código do idioma para OCR
        dpi: Resolução para conversão de páginas
    
    Returns:
        dict com texto e metadados
    """
    content = file.read()
    file.seek(0)
    
    # Obter número de páginas
    doc = fitz.open(stream=content, filetype="pdf")
    total_pages = len(doc)
    doc.close()
    
    # Converter PDF para imagens
    images = pdf_to_images(content, dpi)
    
    # Executar OCR em cada página
    texts = []
    for i, img in enumerate(images):
        page_text = ocr_image(img, lang)
        texts.append({
            'page': i + 1,
            'text': page_text.strip()
        })
    
    # Texto completo
    full_text = '\n\n--- Página {} ---\n\n'.join(['{}'] * len(texts))
    pages_text = []
    for t in texts:
        pages_text.append(f"--- Página {t['page']} ---\n\n{t['text']}")
    
    return {
        'total_pages': total_pages,
        'pages': texts,
        'full_text': '\n\n'.join(pages_text),
        'language': LANGUAGES.get(lang, lang)
    }


def ocr_pdf_to_searchable_pdf(file, lang='por', dpi=300):
    """
    Converte um PDF escaneado em um PDF pesquisável (com texto invisível).
    Usa pytesseract para gerar PDF com camada de texto OCR.
    
    Args:
        file: Arquivo PDF
        lang: Código do idioma para OCR
        dpi: Resolução para OCR
    
    Returns:
        tuple (BytesIO do PDF, nome do arquivo, mimetype)
    """
    content = file.read()
    file.seek(0)
    
    filename = file.filename
    base_name = os.path.splitext(filename)[0]
    
    # Converter PDF para imagens
    images = pdf_to_images(content, dpi)
    
    # Gerar PDFs com OCR para cada página usando pytesseract
    pdf_pages = []
    for img in images:
        # pytesseract gera PDF com texto pesquisável embutido
        pdf_bytes = pytesseract.image_to_pdf_or_hocr(img, lang=lang, extension='pdf')
        pdf_pages.append(pdf_bytes)
    
    # Combinar todas as páginas em um único PDF
    output_doc = fitz.open()
    
    for pdf_bytes in pdf_pages:
        # Abrir cada página OCR
        page_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        output_doc.insert_pdf(page_doc)
        page_doc.close()
    
    # Salvar PDF final
    output = BytesIO()
    output_doc.save(output, garbage=4, deflate=True)
    output_doc.close()
    
    output.seek(0)
    return output, f"{base_name}_pesquisavel.pdf", 'application/pdf'


def ocr_pdf_to_txt(file, lang='por', dpi=300):
    """
    Extrai texto de um PDF e retorna como arquivo TXT.
    
    Args:
        file: Arquivo PDF
        lang: Código do idioma para OCR
        dpi: Resolução para OCR
    
    Returns:
        tuple (BytesIO do TXT, nome do arquivo, mimetype)
    """
    result = ocr_pdf_to_text(file, lang, dpi)
    
    filename = file.filename
    base_name = os.path.splitext(filename)[0]
    
    output = BytesIO()
    output.write(result['full_text'].encode('utf-8'))
    output.seek(0)
    
    return output, f"{base_name}_texto.txt", 'text/plain'


def ocr_pdf_to_docx(file, lang='por', dpi=300):
    """
    Extrai texto de um PDF e retorna como arquivo DOCX.
    
    Args:
        file: Arquivo PDF
        lang: Código do idioma para OCR
        dpi: Resolução para OCR
    
    Returns:
        tuple (BytesIO do DOCX, nome do arquivo, mimetype)
    """
    from docx import Document
    from docx.shared import Pt
    
    result = ocr_pdf_to_text(file, lang, dpi)
    
    filename = file.filename
    base_name = os.path.splitext(filename)[0]
    
    # Criar documento Word
    doc = Document()
    
    for page_data in result['pages']:
        # Adicionar cabeçalho da página
        heading = doc.add_heading(f'Página {page_data["page"]}', level=1)
        
        # Adicionar texto
        para = doc.add_paragraph(page_data['text'])
        para.style.font.size = Pt(11)
        
        # Adicionar quebra de página (exceto na última)
        if page_data['page'] < result['total_pages']:
            doc.add_page_break()
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output, f"{base_name}_texto.docx", 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def get_pdf_preview_for_ocr(content, page_number=0, zoom=0.3):
    """
    Gera preview de uma página do PDF para OCR.
    
    Args:
        content: Conteúdo binário do PDF
        page_number: Número da página (0-indexed)
        zoom: Fator de zoom para preview
    
    Returns:
        Base64 da imagem preview
    """
    import base64
    
    doc = fitz.open(stream=content, filetype="pdf")
    
    if page_number >= len(doc):
        page_number = len(doc) - 1
    
    page = doc.load_page(page_number)
    matrix = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=matrix)
    
    # Converter para PNG
    img_data = pix.tobytes("png")
    doc.close()
    
    return base64.b64encode(img_data).decode('utf-8')


def get_all_pdf_previews_for_ocr(content, zoom=0.3):
    """
    Gera previews de todas as páginas do PDF.
    
    Args:
        content: Conteúdo binário do PDF
        zoom: Fator de zoom para preview
    
    Returns:
        Lista de base64 das imagens preview
    """
    import base64
    
    doc = fitz.open(stream=content, filetype="pdf")
    previews = []
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        matrix = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=matrix)
        img_data = pix.tobytes("png")
        previews.append(base64.b64encode(img_data).decode('utf-8'))
    
    doc.close()
    return previews
